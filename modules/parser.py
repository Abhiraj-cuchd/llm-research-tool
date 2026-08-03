import io
import re
import logging
from pathlib import Path
from docx import Document

from utils.validators import InterviewJSON

logger = logging.getLogger(__name__)


class ParseError(Exception):
    pass


SPEAKER_PATTERNS = [
    (re.compile(r"^Q\s*[:.]\s*", re.IGNORECASE), re.compile(r"^A\s*[:.]\s*", re.IGNORECASE)),
    (re.compile(r"^Interviewer\s*[:.]\s*", re.IGNORECASE), re.compile(r"^Participant\s*[:.]\s*", re.IGNORECASE)),
    (re.compile(r"^\[Q\]\s*", re.IGNORECASE), re.compile(r"^\[A\]\s*", re.IGNORECASE)),
]


def _clean_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _detect_speaker(paragraph, q_pattern, a_pattern):
    text = paragraph.text.strip()
    if not text:
        return None
    if q_pattern.match(text):
        return "Q"
    if a_pattern.match(text):
        return "A"
    return None


def _strip_marker(text: str, pattern) -> str:
    return pattern.sub("", text).strip()


def _bold_is_question(doc) -> list[str]:
    paragraphs = doc.paragraphs
    turns = []
    for para in paragraphs:
        text = _clean_text(para.text)
        if not text:
            continue
        is_bold = any(run.bold for run in para.runs if run.text.strip())
        if is_bold:
            turns.append(("Q", text))
        else:
            turns.append(("A", text))
    return turns


def _alternating_paragraphs(doc) -> list[tuple[str, str]]:
    paragraphs = doc.paragraphs
    cleaned = [_clean_text(p.text) for p in paragraphs if _clean_text(p.text)]
    turns = []
    for i, text in enumerate(cleaned):
        label = "Q" if i % 2 == 0 else "A"
        turns.append((label, text))
    return turns


def parse_docx(filename: str, docx_bytes: bytes) -> InterviewJSON:
    stem = Path(filename).stem
    participant_id = stem

    try:
        doc = Document(io.BytesIO(docx_bytes))
    except Exception as e:
        raise ParseError(f"Cannot read DOCX '{filename}': {e}")

    paragraphs = doc.paragraphs
    if not paragraphs:
        raise ParseError(f"DOCX '{filename}' is empty or has no readable paragraphs.")

    turns = None

    for q_pat, a_pat in SPEAKER_PATTERNS:
        detected = [_detect_speaker(p, q_pat, a_pat) for p in paragraphs]
        labels = [d for d in detected if d is not None]
        if len(labels) >= 2 and any(l == "Q" for l in labels) and any(l == "A" for l in labels):
            qt, at = q_pat, a_pat
            turns = []
            for para in paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                label = _detect_speaker(para, qt, at)
                if label == "Q":
                    turns.append(("Q", _strip_marker(text, qt)))
                elif label == "A":
                    turns.append(("A", _strip_marker(text, at)))
            break

    if turns is None and _bold_is_question(doc):
        turns = _bold_is_question(doc)

    if turns is None:
        turns = _alternating_paragraphs(doc)

    questions = []
    i = 0
    while i < len(turns):
        if turns[i][0] == "Q":
            q_text = turns[i][1]
            answers = []
            i += 1
            while i < len(turns) and turns[i][0] == "A":
                answers.append(turns[i][1])
                i += 1
            combined_answer = " ".join(answers)
            if combined_answer:
                questions.append({"question": q_text, "answer": combined_answer})
            else:
                i += 1
        else:
            combined_answer = []
            while i < len(turns) and turns[i][0] == "A":
                combined_answer.append(turns[i][1])
                i += 1
            if combined_answer and questions:
                questions[-1]["answer"] += " " + " ".join(combined_answer)
            else:
                i += 1

    if not questions:
        all_text = " ".join(_clean_text(p.text) for p in paragraphs if _clean_text(p.text))
        questions = [{"question": "Full transcript", "answer": all_text}]

    return InterviewJSON(participant=participant_id, questions=questions)
